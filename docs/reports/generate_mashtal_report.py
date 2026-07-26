# -*- coding: utf-8 -*-
"""Generate the complete MASHTAL IN448 project report (Word .docx)."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

HERE = Path(__file__).resolve().parent
OUT = HERE / "MASHTAL-IN448-Full-Report.docx"
USECASE_IMG = HERE / "Mashtal-UseCase-Diagram.png"
CLASS_IMG = HERE / "Mashtal-Class-Diagram.png"

TEAM = ["Mohammad Mantach", "Hadi Hojeij", "Hussein Hussein"]
SUPERVISOR = "Dr. Ali Choumane"


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, style=None, size=11, bold=False, italic=False, space_after=8, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def caption(doc, text):
    return add_para(doc, text, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def add_figure(doc, path, width_in, caption_text):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_in))
        caption(doc, caption_text)
    else:
        add_para(doc, f"[Figure missing: {path.name}]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        caption(doc, caption_text)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ========== COVER ==========
    add_para(doc, "Lebanese University", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Faculty of Science — Course IN448", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(
        doc,
        "MASHTAL — Intelligent Agriculture Platform",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )
    add_para(
        doc,
        "A Smart Agriculture Platform Integrating Social Networking, E-Commerce, "
        "and AI-Powered Advisory Support",
        size=12,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=28,
    )
    add_para(doc, "Project submitted in the context of the course IN448", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Prepared by", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for name in TEAM:
        add_para(doc, name, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "", space_after=12)
    add_para(doc, "Supervised by", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, SUPERVISOR, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "Academic Year 2025–2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)

    # ========== ACKNOWLEDGEMENTS ==========
    heading(doc, "Acknowledgements", 1)
    add_para(
        doc,
        "This project was carried out as part of the IN448 course requirements at the Lebanese University. "
        f"The authors express their sincere gratitude to {SUPERVISOR} for his guidance and continuous support "
        "throughout the specification, design, and implementation phases. The team also acknowledges the "
        "agricultural community needs in Lebanon that motivated the problem statement and product vision of MASHTAL.",
    )

    # ========== ABSTRACT ==========
    heading(doc, "Abstract", 1)
    add_para(
        doc,
        "This report presents the complete specification, design, planning, and implementation of "
        "MASHTAL (Intelligent Agriculture Platform), a web-based smart agriculture system that connects "
        "farmers and community users, agricultural businesses, and platform administrators. MASHTAL was "
        "conceived to reduce fragmentation in the agricultural sector by unifying knowledge sharing, "
        "direct communication, online ordering, and intelligent farming support in one platform. "
        "Functionally, the system provides bilingual (English/Arabic) social feeds (posts and discussion "
        "threads), follow and save mechanisms, product marketplace with cart and Stripe checkout, business "
        "subscription (Stripe and Whish Money), seller analytics dashboards, real-time messaging with "
        "Mashtal Support, notifications, content moderation, business reporting, and an AI agricultural "
        "assistant capable of agronomy Q&A and plant-photo disease analysis. Technically, MASHTAL is "
        "implemented as a React/Vite single-page application, an Express/Node.js REST API with MongoDB "
        "persistence, Socket.IO real-time communication, Stripe payment processing, and a Python-based "
        "plant-disease inference service. This document covers the business domain, requirements analysis "
        "and specification (SRS-style), UML use-case and class design (with diagrams), architecture, "
        "project planning, technologies, implementation, testing, teamwork, and conclusions.",
    )
    add_para(
        doc,
        "Keywords: smart agriculture; e-commerce; social networking; AI advisory; plant disease detection; "
        "web platform; software engineering; requirements engineering; UML.",
        italic=True,
    )

    page_break(doc)

    # ========== TOC ==========
    heading(doc, "Table of Contents", 1)
    add_para(
        doc,
        "In Microsoft Word: place the cursor below, then use References → Table of Contents → Automatic Table "
        "to refresh page numbers. Headings in this document use Word heading styles.",
        italic=True,
    )
    for item in [
        "1. Introduction",
        "2. Requirement Gathering, Analysis, and Specification",
        "3. High-Level Design Specification",
        "4. Application Conception",
        "5. Development",
        "6. Conclusion",
        "Bibliography",
        "Appendices",
    ]:
        add_para(doc, item, space_after=4)

    heading(doc, "Table of Figures", 1)
    for f in [
        "Figure 1. MASHTAL use-case diagram (UML)",
        "Figure 2. High-level layered architecture of MASHTAL",
        "Figure 3. MASHTAL domain class diagram (UML)",
        "Figure 4. Conceptual stakeholder and value-flow overview",
    ]:
        add_para(doc, f, size=10, space_after=2)

    heading(doc, "Table of Tables", 1)
    for t in [
        "Table 1. Comparison of existing similar platforms",
        "Table 2. User classes and characteristics",
        "Table 3. MoSCoW prioritization of major requirements",
        "Table 4. Functional requirements catalog (REQ)",
        "Table 5. Project planning phases and milestones",
        "Table 6. Technology stack summary",
        "Table 7. MongoDB domain models inventory",
        "Table 8. Selected use-case textual descriptions",
    ]:
        add_para(doc, t, size=10, space_after=2)

    page_break(doc)

    # ========== CHAPTER 1 ==========
    heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "Agriculture remains a strategic sector for food security, rural livelihoods, and local economies. "
        "Yet farmers, agricultural suppliers, and community practitioners often face fragmented digital tools: "
        "social media for advice that is not specialized, marketplaces that are not agriculture-focused, and "
        "advisory channels that are offline or hard to reach. MASHTAL was designed as an Intelligent Agriculture "
        "Platform that connects these stakeholders in one coherent web system.",
    )

    heading(doc, "1.1 The Business Domain", 2)
    add_para(
        doc,
        "The business domain is digital agriculture services with three intertwined value streams: (1) knowledge "
        "and community (posts, threads, comments, follows, translation); (2) commerce (agricultural products, "
        "cart, payments, orders, reviews); and (3) intelligence and support (AI plant diagnosis and agronomy "
        "chat, real-time messaging, platform support, moderation). The primary geographic focus is Lebanon "
        "(locations, language needs, and local payment options such as Whish Money alongside international "
        "card payments via Stripe).",
    )
    add_para(
        doc,
        "Stakeholders include community users and farmers seeking advice and supplies; agricultural businesses "
        "seeking online presence, sales, and analytics; and platform administrators ensuring trust, payments "
        "integrity, and support quality. External systems (Stripe, Google OAuth, email SMTP, Whish Money, "
        "AI inference, translation, and location search) complete the ecosystem.",
    )

    heading(doc, "1.2 About the Modeled Application (Purpose, User Personas, and Product Perspective)", 2)
    add_para(
        doc,
        "Original product vision: MASHTAL is a smart agriculture platform that connects farmers, agricultural "
        "businesses, engineers, and users, enabling knowledge sharing, direct communication, online ordering, "
        "and intelligent farming support to enhance productivity and collaboration in the agricultural sector. "
        "It integrates AI-powered agricultural assistance for problem diagnosis and advisory support with "
        "social networking and e-commerce.",
    )
    add_para(
        doc,
        "In the implemented system, this vision is realized through a clear role model aligned with the UML "
        "use-case diagram: Guest (unauthenticated), User (authenticated visitor / community member — including "
        "farmers and practitioners who share knowledge and buy products), Business (seller with subscription), "
        "and Admin (platform operator). Agricultural engineers participate as Users in community discussions "
        "and messaging; businesses provide commercial products and content; AI complements human advisory "
        "support for plant problems.",
    )
    add_para(doc, "Primary personas:")
    bullet(doc, "Community User / Farmer — browses content, asks AI, buys supplies, chats, follows businesses.")
    bullet(doc, "Agricultural Business — pays subscription, lists products, manages orders, posts content, views analytics.")
    bullet(doc, "Administrator — manages users, subscriptions, transactions, orders, reports, and support chat.")
    bullet(doc, "Guest — discovers the platform, searches, uses AI, then registers/signs in.")

    heading(doc, "1.3 Analysis of Existing Similar Applications", 2)
    add_para(
        doc,
        "Comparable solutions typically specialize in one dimension (social feed, marketplace, or advisory). "
        "MASHTAL’s differentiation is the integrated combination of agri-focused social networking, verified "
        "seller commerce with subscription control, realtime chat/support, and AI plant advisory in one product.",
    )
    caption(doc, "Table 1. Comparison of existing similar platforms")
    add_table(
        doc,
        [
            "Platform type / example",
            "Agri social feed",
            "Verified seller shop + checkout",
            "Realtime chat",
            "AI plant / farm advisory",
            "Seller analytics + subscription",
        ],
        [
            ["General social networks", "Partial", "No / weak", "Yes", "No", "No"],
            ["General marketplaces", "No", "Yes", "Partial", "No", "Partial"],
            ["Agri advice apps / chatbots", "Weak", "No", "Partial", "Yes / partial", "No"],
            ["Local farm directories", "No", "Partial", "No", "No", "Partial"],
            ["MASHTAL (this project)", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ],
    )

    heading(doc, "1.4 Plan of the Document", 2)
    add_para(
        doc,
        "Chapter 2 presents requirements gathering, analysis, and specification (including features, NFRs, "
        "MoSCoW, and use cases). Chapter 3 provides the high-level design specification (security, UI, "
        "architecture, interfaces, database overview, and planning). Chapter 4 details application conception "
        "(database and UML class model). Chapter 5 describes development approach, challenges, testing, "
        "deployment, teamwork, and reflection. Chapter 6 concludes and outlines future work. Appendices "
        "reference diagrams and implementation evidence.",
    )

    page_break(doc)

    # ========== CHAPTER 2 ==========
    heading(doc, "2. Requirement Gathering, Analysis, and Specification", 1)

    heading(doc, "2.1 Introduction", 2)
    heading(doc, "2.1.1 Purpose", 3)
    add_para(
        doc,
        "This chapter defines what MASHTAL must do and under which constraints. It serves as the Software "
        "Requirements Specification (SRS) baseline for design and implementation, and as evidence of "
        "requirements engineering for the IN448 evaluation.",
    )
    heading(doc, "2.1.2 Document Conventions", 3)
    bullet(doc, "REQ-F-xxx: functional requirement.")
    bullet(doc, "REQ-N-xxx: non-functional requirement.")
    bullet(doc, "UC-*: use-case identifiers grouped by actor (Guest, User/Visitor, Business, Admin, External).")
    bullet(doc, "Actors in UML: Guest, User, Business, Admin (inheritance: Business and Admin inherit User; User inherits Guest capabilities where applicable).")
    bullet(doc, "Implementation role names: guest (unauthenticated), visitor, business, admin.")

    heading(doc, "2.1.3 Intended Audience and Reading Suggestions", 3)
    add_para(
        doc,
        "Intended for the IN448 instructor, the development team, and future maintainers. Readers new to the "
        "project should read Sections 1.1–1.2, then 2.2–2.5, then the use-case diagram and selected textual "
        "use cases. Designers/developers should focus on Chapters 3–5.",
    )

    heading(doc, "2.1.4 Product Scope", 3)
    add_para(
        doc,
        "In scope: bilingual web SPA; authentication (email/password, Google OAuth, email verification); "
        "social posts and threads with comments, likes, shares, saves, follows; marketplace catalog, cart, "
        "Stripe checkout, order lifecycle; business subscription (Stripe + Whish); business dashboard "
        "analytics; admin console; realtime chat and Mashtal Support; notifications; AI assistant "
        "(text + plant photo); translation; Lebanon location search; content moderation; business reports. "
        "Out of scope for this delivery: native mobile apps, full logistics/fleet tracking, multi-country "
        "tax engines, and formal expert booking marketplace as a separate paid role.",
    )

    heading(doc, "2.1.5 References", 3)
    bullet(doc, "IEEE 830 / modern SRS practices for requirements structure.")
    bullet(doc, "UML 2.x for use-case and class modeling.")
    bullet(doc, "Stripe, Socket.IO, MongoDB, React, Express official documentation.")
    bullet(doc, "Project UML artifacts and implemented codebase (client + server).")

    heading(doc, "2.2 Overall Description", 2)
    heading(doc, "2.2.1 Product Perspective", 3)
    add_para(
        doc,
        "MASHTAL is a new standalone web platform. It is not an extension of a legacy monolith. It integrates "
        "external services for payments (Stripe, Whish), identity (Google OAuth), email verification, AI "
        "inference, translation, and geolocation search, while owning the core domain data in MongoDB.",
    )

    heading(doc, "2.2.2 Product Functions", 3)
    add_para(doc, "Major functional areas of MASHTAL:")
    numbered(doc, "Authentication and profiles (register, sign-in, Google OAuth, email verify, profile/avatar/cover).")
    numbered(doc, "Discover / Home, Posts, Threads (create, engage, filter, translate, share).")
    numbered(doc, "Unified search and filters (businesses, products, posts, threads).")
    numbered(doc, "Follow graph and Saved items.")
    numbered(doc, "Marketplace: products by category, cart, Stripe checkout, purchase history, cancel/ready/complete.")
    numbered(doc, "Ratings and reviews (products and businesses); business reporting.")
    numbered(doc, "Business onboarding and paid subscription; seller tools (products, orders, analytics).")
    numbered(doc, "Realtime messaging, block/unblock, Mashtal Support with admin locks.")
    numbered(doc, "Notifications (follows, likes, comments, orders, chat, subscriptions, admin actions).")
    numbered(doc, "AI agricultural assistant (agronomy Q&A + plant-photo disease analysis + product tips).")
    numbered(doc, "Administration (users, businesses, subscriptions, transactions, orders, reports, Whish verify).")
    numbered(doc, "Internationalization EN/AR with RTL and preferred language persistence.")

    heading(doc, "2.2.3 User Classes and Characteristics", 3)
    caption(doc, "Table 2. User classes and characteristics")
    add_table(
        doc,
        ["User class", "Access", "Key goals", "Technical expectation"],
        [
            ["Guest", "Public pages", "Discover, search, try AI, register", "Low friction; clear CTAs to sign in"],
            ["User (Visitor)", "Authenticated", "Social engage, buy, chat, notify", "Secure session; responsive UX"],
            ["Business (inactive)", "Authenticated + role", "Renew plan; edit profile; limited content", "Clear renew path; products hidden"],
            ["Business (active)", "Paid subscription", "Sell, fulfill, analytics, post", "Dashboard KPIs; order actions"],
            ["Admin", "Elevated", "Trust, finance, support, moderation", "Role-gated APIs; auditability"],
        ],
    )

    heading(doc, "2.2.4 Operating Environment", 3)
    bullet(doc, "Client: modern desktop/mobile browsers (Chromium, Firefox, Safari/Edge).")
    bullet(doc, "Server: Node.js runtime hosting Express API and Socket.IO.")
    bullet(doc, "Database: MongoDB.")
    bullet(doc, "AI sidecar: Python FastAPI service for plant-disease classification.")
    bullet(doc, "Network: HTTPS in production; HTTP localhost for development.")

    heading(doc, "2.2.5 Design and Implementation Constraints", 3)
    bullet(doc, "Must support English and Arabic (RTL).")
    bullet(doc, "Payments must use Stripe for card checkout; Whish for alternative subscription payment.")
    bullet(doc, "Business selling capabilities require an active subscription.")
    bullet(doc, "Admin support replies use a shared Mashtal Support identity with concurrency locks.")
    bullet(doc, "Sensitive secrets remain in environment variables, never in source control.")

    heading(doc, "2.2.6 User Documentation", 3)
    add_para(
        doc,
        "In-app UI labels, validation messages, and bilingual strings act as primary user guidance. "
        "About / Privacy / Terms / Cookies pages provide legal and informational content. This report "
        "and UML diagrams serve as technical documentation for evaluators and maintainers.",
    )

    heading(doc, "2.2.7 Assumptions and Dependencies", 3)
    bullet(doc, "Users have internet access and a compatible browser.")
    bullet(doc, "Stripe test/live keys and webhook connectivity are configured for payments.")
    bullet(doc, "SMTP credentials exist for email verification.")
    bullet(doc, "AI service and optional Hugging Face tokens are available for advisory quality.")
    bullet(doc, "Lebanon location dataset / Nominatim access is available for place search.")

    heading(doc, "2.3 External Interface Requirements", 2)
    heading(doc, "2.3.1 User Interfaces", 3)
    add_para(
        doc,
        "The UI is a responsive SPA with navigation (Home, Posts, Threads, Shop, Search, Businesses, "
        "Dashboard/Admin, Chats, Notifications, Profile). Key interactive surfaces: feeds with filters and "
        "side panels; product cards/detail modal; cart and checkout with Stripe Elements; business and admin "
        "dashboards with charts; floating AI assistant; chat inbox with presence; bilingual toggle.",
    )
    heading(doc, "2.3.2 Hardware Interfaces", 3)
    add_para(doc, "No specialized hardware. Camera/file upload for plant photos and media; standard input devices.")
    heading(doc, "2.3.3 Software Interfaces", 3)
    bullet(doc, "REST JSON API between React client and Express server.")
    bullet(doc, "Socket.IO for realtime chat, typing/presence, and support locks.")
    bullet(doc, "Stripe PaymentIntents + webhooks.")
    bullet(doc, "Google OAuth token verification.")
    bullet(doc, "SMTP email (Nodemailer).")
    bullet(doc, "Whish Money subscription submission/verification path.")
    bullet(doc, "Python AI disease service HTTP API.")
    bullet(doc, "Translation endpoint EN↔AR.")
    bullet(doc, "Location search (Lebanon-focused Nominatim proxy / local dataset).")
    heading(doc, "2.3.4 Communications Interfaces", 3)
    add_para(
        doc,
        "HTTPS/HTTP for REST; WebSocket (Socket.IO) for realtime; webhook callbacks from Stripe (and Whish "
        "verification flows). CORS configured between frontend origin and API origin.",
    )

    heading(doc, "2.4 Requirements Gathering", 2)
    add_para(
        doc,
        "Requirements were gathered through: (1) domain problem analysis of fragmented agri digital tools; "
        "(2) stakeholder personas (farmer/user, business, admin); (3) competitive comparison (Table 1); "
        "(4) iterative workshops refining must-have flows (auth, social, shop, subscription, chat, AI, admin); "
        "(5) continuous validation against the implemented prototype. Requirements were classified by actor "
        "and prioritized with MoSCoW.",
    )

    heading(doc, "2.5 System Features", 2)

    heading(doc, "2.5.1 Authentication and Profiles", 3)
    add_para(
        doc,
        "REQ-F-001: Guests can register as User (visitor) or start business registration; sign in with email/"
        "password or Google OAuth; verify email. Authenticated users manage profile fields, avatar, cover, "
        "phone, location, bio, and preferred language. Users can convert to business (pending profile then "
        "subscription activation). Sessions use JWT; /auth/me restores identity.",
    )

    heading(doc, "2.5.2 Discover, Posts, Threads, Search, Follow, and Save", 3)
    add_para(
        doc,
        "REQ-F-010: Discover/Home shows featured/trusted businesses and combined feeds. Posts (typically image "
        "content by business/admin) and Threads (discussion) support create/edit/delete (authorized roles), "
        "likes, nested comments/replies, shares (external and in-app), and saves. Before a post or thread "
        "is published, automated content moderation checks text and images: sexual content, weapons/guns, "
        "and other policy-violating material are rejected and the content is not stored or shown. The same "
        "safety checks apply to comments, chat messages, product text/images, and profile media where "
        "configured. Feeds support All / Following / Businesses filters. Unified Search finds businesses, "
        "products, posts, and threads. Follow/save features are supported.",
    )

    heading(doc, "2.5.3 Marketplace, Cart, Checkout, Orders, Reviews, and Reports", 3)
    add_para(
        doc,
        "REQ-F-020: Shop lists products by categories (seeds, tools, fertilizers, plants, irrigation, "
        "equipment, trees, medicament, other) with search, filters, sorting, and grid/list views. Users "
        "manage cart, checkout with shipping data, and pay via Stripe PaymentIntent. Orders progress through "
        "processing → ready → completed / cancelled; buyers can cancel under policy; businesses mark ready; "
        "admins can override status. Product and business reviews/ratings are supported. Users can report "
        "businesses for admin review.",
    )

    heading(doc, "2.5.4 Business Subscription and Seller Tools", 3)
    add_para(
        doc,
        "REQ-F-030: Businesses activate/renew subscription via Stripe card or Whish Money transfer (admin "
        "verification path). Active subscription enables product visibility, order acceptance, and analytics. "
        "Inactive subscription hides products from public shop and restricts seller ops to renew, profile, "
        "and basic content management (aligned with the use-case diagram note). Business dashboard provides "
        "KPIs, revenue/sales charts, product performance, and order lists. Business profile includes hours, "
        "specialties, about, website, and Whish payout identity.",
    )

    heading(doc, "2.5.5 Messaging and Notifications", 3)
    add_para(
        doc,
        "REQ-F-040: Realtime 1:1 messaging with Socket.IO; edit/delete own messages (policy window); "
        "block/unblock; presence indicators; open Mashtal Support. Admins share a support inbox with "
        "acquire/renew/release locks to avoid concurrent reply conflicts. Notifications cover follows, "
        "likes, comments, reviews, orders, chat aggregates, subscription events, payments, reports, and "
        "admin warnings, with mark-read and clear actions and deep-links.",
    )

    heading(doc, "2.5.6 AI Agricultural Assistance", 3)
    add_para(
        doc,
        "REQ-F-050: Floating AI assistant available to guests and authenticated users. Supports agronomy "
        "Q&A, casual agricultural chat, plant-photo disease analysis via local Python classifier, treatment "
        "guidance, and product recommendation hints. Off-topic prompts are refused. Content safety "
        "mechanisms (text toxicity and image checks) protect uploads and posts.",
    )

    heading(doc, "2.5.7 Administration and Trust", 3)
    add_para(
        doc,
        "REQ-F-060: Admin dashboard covers overview metrics and Mashtal income (tax + subscription fees), "
        "user CRUD and role/subscription controls, businesses list, subscriptions and notify-expiring, "
        "money transaction ledger (seller/tax/subscription legs), all orders, report queue resolution "
        "(dismiss / warn / stronger action), Whish payment verification, support chat, and admin-authored "
        "posts/threads.",
    )

    heading(doc, "2.5.8 Internationalization and Location", 3)
    add_para(
        doc,
        "REQ-F-070: Full UI bilingual English/Arabic with RTL layout. Preferred language can persist on the "
        "user profile. Location fields and search filters focus on Lebanon places. (Optional UX: an on-demand "
        "“see translation” helper may assist bilingual reading of individual posts/comments; it is a "
        "convenience feature, not a core business capability.)",
    )

    heading(doc, "2.5.9 Content Moderation and Community Safety", 3)
    add_para(
        doc,
        "REQ-F-080: MASHTAL must not publish unsafe content. When a business or admin creates or updates a "
        "post or thread (text and/or image), the server runs automated safety checks. If the image is "
        "classified as sexual/NSFW, weapons/guns, or otherwise disallowed, or if the text fails toxicity/"
        "policy checks, publishing is blocked: the upload is discarded and the client receives a clear "
        "rejection message. Moderated surfaces also include comments, selected chat text, product "
        "descriptions/images, and profile avatar/cover. Manual trust remains available via business reports "
        "handled by admins (Handle reports).",
    )

    heading(doc, "2.6 Other Nonfunctional Requirements", 2)
    heading(doc, "2.6.1 Performance Requirements", 3)
    add_para(
        doc,
        "REQ-N-001: Feed and shop lists should load with pagination/lazy patterns acceptable for interactive "
        "use; chat messages should appear near-realtime over Socket.IO; dashboards should render charts "
        "without blocking the whole SPA.",
    )
    heading(doc, "2.6.2 Safety Requirements", 3)
    add_para(
        doc,
        "REQ-N-010: AI advice is assistive, not a substitute for certified agronomist judgment in critical "
        "cases. Posts, threads, and other user-generated media/text must pass automated moderation before "
        "publication (sexual content, weapons, and harmful text are blocked). Legal pages disclose platform "
        "policies.",
    )
    heading(doc, "2.6.3 Security Requirements", 3)
    add_para(
        doc,
        "REQ-N-020: Passwords hashed (bcrypt); JWT auth; role-based authorization on protected routes; "
        "CORS; secrets in .env; Stripe webhook verification; email verification gate; admin-only routes "
        "for privileged operations; support lock to reduce race conditions.",
    )
    heading(doc, "2.6.4 Software Quality Attributes", 3)
    add_para(
        doc,
        "REQ-N-030: Maintainability via modular controllers/models and typed client API layers; usability "
        "via consistent navigation and bilingual UX; reliability via webhook-driven payment confirmation "
        "and clear order status machine; portability as a standard web stack.",
    )
    heading(doc, "2.6.5 Business Rules", 3)
    bullet(doc, "Only active subscribed businesses can sell (products public + order acceptance).")
    bullet(doc, "Inactive businesses may renew, manage profile, and manage limited content.")
    bullet(doc, "Platform tax/fee legs and subscription fees are recorded in the money ledger.")
    bullet(doc, "Order cancellation and status transitions follow defined server-side rules.")
    bullet(doc, "Admins resolve reports and may warn businesses.")

    heading(doc, "2.7 Requirements Analysis", 2)
    caption(doc, "Table 3. MoSCoW prioritization of major requirements")
    add_table(
        doc,
        ["MoSCoW", "Capabilities"],
        [
            [
                "Must have",
                "Auth + profiles; posts/threads/comments with automated content moderation; shop cart Stripe "
                "checkout; business subscription; orders; chat; notifications; admin basics; AI assistant "
                "core; EN/AR UI",
            ],
            [
                "Should have",
                "Whish subscription path; analytics charts; reports; saved items; search filters; support locks",
            ],
            [
                "Could have",
                "On-demand post translation helper; richer recommendations; deeper logistics; broader "
                "automated tests; native mobile",
            ],
            [
                "Won’t have (this delivery)",
                "Full multi-country tax engine; fleet tracking; standalone paid engineer booking marketplace",
            ],
        ],
    )

    caption(doc, "Table 4. Functional requirements catalog (REQ) — consolidated")
    add_table(
        doc,
        ["ID", "Statement"],
        [
            ["REQ-F-001", "Authenticate via email/password or Google; verify email; manage profile media/fields."],
            ["REQ-F-010", "Social discover/posts/threads with engage, filter, share, follow, save, search."],
            ["REQ-F-020", "Marketplace catalog, cart, Stripe pay, orders, reviews, business reports."],
            ["REQ-F-030", "Business subscription (Stripe/Whish) gates seller products/orders/analytics."],
            ["REQ-F-040", "Realtime chat, block, Mashtal Support locks; notification center."],
            ["REQ-F-050", "AI agronomy Q&A and plant-photo disease analysis with safety checks."],
            ["REQ-F-060", "Admin console for users, businesses, subscriptions, ledger, orders, reports, Whish."],
            ["REQ-F-070", "Bilingual EN/AR RTL + Lebanon location assistance."],
            ["REQ-F-080", "Block unsafe posts/threads/images/text (sexual, weapons, toxic) before publish."],
            ["REQ-N-020", "JWT, hashing, RBAC, webhook verification, secret isolation."],
        ],
    )

    heading(doc, "2.8 Use Cases", 2)
    add_para(
        doc,
        "Figure 1 presents the complete UML use-case diagram of MASHTAL. Actors follow generalization: "
        "Guest ← User ← Business and Guest ← User ← Admin. Include relationships model mandatory steps "
        "(e.g., Verify email within Register/Sign in; Process payment within Checkout and Renew "
        "subscription). Extend relationships model optional specializations (Answer agronomy question / "
        "Analyze plant photo for AI; Share for social content). Creating posts/threads is subject to "
        "automated moderation: disallowed images or text prevent publication.",
    )
    add_figure(doc, USECASE_IMG, 6.5, "Figure 1. MASHTAL PLATFORM — Use Case Diagram (UML)")

    add_para(doc, "Actor summary matching Figure 1:")
    bullet(doc, "Guest: Register/Sign in (includes Verify email), Browse content, Use AI assistant, Search and filter.")
    bullet(
        doc,
        "User: Manage profile; Posts/threads/comments (extend Share; content must pass safety checks); "
        "Follow and save; Realtime messaging; Manage notifications; Ratings and reviews; Manage cart; "
        "Checkout and payment (includes Process payment); Manage purchases.",
    )
    bullet(
        doc,
        "Business: Manage business profile; Manage subscription; Manage products; Manage business orders; "
        "View business analytics; Manage business posts & threads; Chat with customers; Renew subscription "
        "(includes Process payment). Active subscription enables products/orders/analytics; inactive limits "
        "to renew, profile, and basic content.",
    )
    bullet(
        doc,
        "Admin: View admin dashboard; Manage users; Manage businesses; Manage subscriptions; Manage "
        "transactions; Manage all orders; Handle reports; Manage support conversations; Create admin posts & threads.",
    )

    add_para(
        doc,
        "The detailed catalog implemented in the system expands these ovals into fine-grained use cases "
        "(browse shop filters, cancel order, mark ready, Whish verify, support lock acquire/release, etc.). "
        "Appendix A lists the catalog groups for traceability.",
    )

    heading(doc, "2.9 Use Cases Textual Description (Selected)", 2)
    caption(doc, "Table 8. Selected use-case textual descriptions")
    add_table(
        doc,
        ["Use case", "Actor", "Main success scenario (summary)", "Exceptions / rules"],
        [
            [
                "Register / Sign in",
                "Guest",
                "Open form → submit credentials or Google → verify email if required → session created",
                "Invalid credentials; unverified email gate",
            ],
            [
                "Create / update post or thread",
                "Business / Admin",
                "Compose text and optional image → server runs safety checks → if allowed, publish to feed",
                "Rejected if image is sexual/NSFW or weapons/guns, or text violates policy; upload discarded; not published",
            ],
            [
                "Use AI assistant",
                "Guest/User",
                "Open assistant → ask question or upload plant photo → receive advice / diagnosis",
                "Off-topic refusal; low-confidence diagnosis warning",
            ],
            [
                "Checkout and payment",
                "User",
                "Cart → checkout shipping → Stripe PaymentIntent → webhook confirms → order created",
                "Payment failure; empty cart; inactive seller products filtered",
            ],
            [
                "Manage subscription / Renew",
                "Business",
                "Pay Stripe or submit Whish → activation/renewal dates set → seller tools enabled",
                "Whish pending until admin verify; expiry disables selling",
            ],
            [
                "Manage business orders",
                "Business (active)",
                "List orders → inspect items → mark ready → buyer notified",
                "Inactive subscription blocks seller ops",
            ],
            [
                "Handle reports",
                "Admin",
                "Open reports queue → review → dismiss / warn / take action → notify parties",
                "Missing evidence; already resolved",
            ],
            [
                "Manage support conversations",
                "Admin",
                "Open support inbox → acquire lock → reply via shared Support identity → release lock",
                "Lock held by another admin",
            ],
        ],
    )

    page_break(doc)

    # ========== CHAPTER 3 ==========
    heading(doc, "3. High-Level Design Specification", 1)
    add_para(
        doc,
        "This chapter translates requirements into a high-level technical design covering security, hardware, "
        "UI, interfaces, architecture, analytics, database overview, and design-phase planning.",
    )

    heading(doc, "3.1 Security", 2)
    bullet(doc, "Authentication: JWT after email/password (bcrypt) or Google OAuth verification.")
    bullet(doc, "Authorization: role checks (visitor/business/admin) on API routes and UI navigation guards.")
    bullet(doc, "Transport: HTTPS recommended in production; CORS restricted to frontend origin.")
    bullet(doc, "Payments: Stripe server-side PaymentIntents; webhook signature verification.")
    bullet(doc, "Data: secrets in environment variables; least privilege for admin operations.")
    bullet(doc, "Trust & safety: moderation on content/uploads; report workflow; support locks.")

    heading(doc, "3.2 Hardware", 2)
    add_para(
        doc,
        "Development uses standard developer workstations. Production expects a cloud or VPS host for API, "
        "static frontend assets, MongoDB (managed or self-hosted), and optional GPU/CPU host for the Python "
        "AI service. End users need a smartphone or computer with camera for plant photos.",
    )

    heading(doc, "3.3 User Interface", 2)
    add_para(
        doc,
        "UI architecture is component-based React with shared layout (nav/footer), page routes, modals "
        "(product detail, post), dashboards (business/admin), and global widgets (AI FAB, toasts, "
        "scroll-to-top). Design goals: clear agri branding, bilingual RTL, accessible forms, and "
        "role-appropriate menus. Charts (Recharts) visualize seller and admin analytics.",
    )

    heading(doc, "3.4 Internal Interfaces", 2)
    add_para(
        doc,
        "Internal interfaces are REST resource modules (/api/auth, /users, /businesses, /posts, /threads, "
        "/comments, /products, /orders, /saved, /notifications, /chat, /reviews, /reports, /dashboard, "
        "/admin, /payments/..., /ai, /translate, /locations, /health) plus Socket.IO events for chat/"
        "presence/support. Client API modules and shared TypeScript types mirror server contracts.",
    )

    heading(doc, "3.5 External Interfaces", 2)
    bullet(doc, "Stripe — checkout PaymentIntents, subscription PaymentIntents, webhooks.")
    bullet(doc, "Google — OAuth ID token verification for sign-in/up.")
    bullet(doc, "SMTP — verification emails.")
    bullet(doc, "Whish Money — subscription transfer submit + admin verify.")
    bullet(doc, "AI Disease Service — image classification HTTP calls.")
    bullet(doc, "Translation service — EN↔AR.")
    bullet(doc, "Location search — Lebanon Nominatim/local dataset.")
    bullet(doc, "External social share URL targets (WhatsApp, X, Facebook, LinkedIn, Telegram, Email).")

    heading(doc, "3.6 Architecture", 2)
    add_para(
        doc,
        "MASHTAL follows a layered client–server architecture: Presentation (React SPA) → API Gateway/"
        "Controllers (Express) → Domain Services/Utils (orders, subscription, presence, notifications, "
        "moderation) → Persistence (MongoDB/Mongoose) → External adapters (Stripe, Google, SMTP, AI, "
        "Whish). Realtime is a parallel channel via Socket.IO sharing auth context.",
    )
    add_para(doc, "Logical layers:")
    numbered(doc, "Client presentation and state (pages, contexts, API clients, i18n).")
    numbered(doc, "HTTP/WebSocket edge (Express + Socket.IO + middleware auth).")
    numbered(doc, "Application controllers and business rules.")
    numbered(doc, "Data models and indexes in MongoDB.")
    numbered(doc, "External systems and AI sidecar.")
    caption(doc, "Figure 2. High-level layered architecture of MASHTAL (textual model)")
    add_para(
        doc,
        "[Browser SPA] ↔ HTTPS/REST + Socket.IO ↔ [Node/Express API] ↔ [MongoDB]\n"
        "                                         ↕\n"
        "                    [Stripe | Google | SMTP | Whish | Translate | Locations]\n"
        "                                         ↕\n"
        "                              [Python AI Disease Service]",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    heading(doc, "3.7 Reports / Analytics", 2)
    add_para(
        doc,
        "Business analytics: revenue, units sold, growth KPIs; revenue-by-product; sales over week/month/year; "
        "product performance table; order management views. Admin analytics: user/business counts, active "
        "paid subscriptions, expiring soon, Mashtal income totals (tax + subscription), transaction ledger "
        "inspection, order and report queues.",
    )

    heading(doc, "3.8 Database", 2)
    add_para(
        doc,
        "MongoDB document store with Mongoose schemas. User is the central aggregate with role and optional "
        "embedded businessProfile / pendingBusinessProfile. Commerce uses Product, Order (embedded items/"
        "shipping), Payment, SubscriptionPayment, MoneyTransaction. Social uses Post, Thread, Comment, "
        "Follow, SavedItem. Trust uses Review, BusinessReview, BusinessReport. Messaging uses Conversation "
        "and ChatMessage. Notification stores polymorphic alerts. Detailed class design appears in Chapter 4.",
    )

    heading(doc, "3.9 Project Planning (Design Phase Planning)", 2)
    caption(doc, "Table 5. Project planning phases and milestones")
    add_table(
        doc,
        ["Phase", "Main activities", "Key outputs"],
        [
            ["1. Inception", "Vision, personas, competitor scan", "Problem statement; product vision"],
            ["2. Requirements", "Elicitation, MoSCoW, use cases", "SRS chapter; use-case diagram"],
            ["3. Design", "Architecture, class model, UI map", "HLD; class diagram; API outline"],
            ["4. Implementation", "Vertical slices: auth→social→shop→pay→chat→AI→admin", "Working SPA + API"],
            ["5. Verification", "Scenario tests, Stripe test mode, bugfix", "Stable demo build"],
            ["6. Delivery", "Report, diagrams, presentation", "IN448 submission package"],
        ],
    )
    add_para(
        doc,
        "Methodology: iterative Agile-style increments with frequent integration demos among team members. "
        "Risk items (payments, realtime concurrency, bilingual RTL, AI practicality) were tackled early with "
        "explicit technical spikes.",
    )

    heading(doc, "3.10 Conclusion of High-Level Design", 2)
    add_para(
        doc,
        "The high-level design establishes a secure, modular, integration-friendly architecture capable of "
        "delivering the social, commercial, administrative, and AI features required by Chapter 2, while "
        "remaining implementable within the academic project timeframe.",
    )

    page_break(doc)

    # ========== CHAPTER 4 ==========
    heading(doc, "4. Application Conception", 1)

    heading(doc, "4.1 Introduction", 2)
    add_para(
        doc,
        "Application conception details how domain concepts become persistent structures and object "
        "relationships. MASHTAL uses a document-oriented model that still exhibits clear UML class "
        "associations and multiplicities.",
    )

    heading(doc, "4.2 Database", 2)
    caption(doc, "Table 7. MongoDB domain models inventory")
    add_table(
        doc,
        ["Model", "Responsibility"],
        [
            ["User", "Identity, roles, follow/block lists, subscription fields, embedded business profile"],
            ["Product", "Seller catalog item: price, stock, category, images, ratings"],
            ["Order", "Buyer order with embedded items, shipping, status, payment link"],
            ["Payment", "Stripe checkout record with cart snapshot and money legs"],
            ["SubscriptionPayment", "Business plan payment (Stripe or Whish transfer metadata)"],
            ["MoneyTransaction", "Ledger entries (seller, tax, subscription)"],
            ["Post / Thread", "Social content authored by business/admin (and engagement counters)"],
            ["Comment", "Polymorphic comments/replies on posts or threads"],
            ["Review / BusinessReview", "Product and business ratings"],
            ["BusinessReport", "User reports and admin resolution fields"],
            ["Follow", "Explicit follower→following edges"],
            ["Conversation / ChatMessage", "DM and support threads; optional shared post preview"],
            ["Notification", "Typed alerts with read state and optional aggregates"],
            ["SavedItem", "Polymorphic saved post/thread/product references"],
        ],
    )
    add_para(
        doc,
        "Key rules: Order embeds OrderItems and Shipping; Payment embeds cart items and legs for split "
        "accounting; Conversation may be marked isSupport with lock fields for admin concurrency; "
        "subscriptionStatus/dates on User gate seller capabilities.",
    )

    heading(doc, "4.3 UML Class Diagram", 2)
    add_para(
        doc,
        "Figure 3 shows the domain class diagram for MASHTAL. User is the central class associated with "
        "products sold, orders placed, payments, subscriptions, social authorship, reviews, reports, "
        "follows, conversations, messages, notifications, and saved items. Composition is used for "
        "embedded structures (order items, shipping, payment legs, business hours). Comments attach "
        "polymorphically to Post or Thread. The diagram provides the structural blueprint corresponding "
        "to Mongoose models in the implementation.",
    )
    add_figure(doc, CLASS_IMG, 6.5, "Figure 3. MASHTAL — Domain Class Diagram (UML)")

    add_para(doc, "Relationship highlights:")
    bullet(doc, "User 1—* Product (sells); User 1—* Order (places); Order *— OrderItem; OrderItem → Product.")
    bullet(doc, "User 1—* Post/Thread/Comment; Comment 0..1—* Comment (replies).")
    bullet(doc, "User *—* User via followers/following/blocked and Follow entity.")
    bullet(doc, "Conversation 1—* ChatMessage; participants are Users; support lock references an Admin user.")
    bullet(doc, "Payment / SubscriptionPayment / MoneyTransaction track commercial and subscription money flows.")
    bullet(doc, "Review → Product; BusinessReview/BusinessReport → business User.")

    heading(doc, "4.4 Conclusion", 2)
    add_para(
        doc,
        "The conception phase yields a coherent domain model that supports social, commerce, messaging, "
        "trust, and subscription concerns without forcing an overly rigid relational schema, while remaining "
        "expressible as a standard UML class diagram for academic evaluation.",
    )

    page_break(doc)

    # ========== CHAPTER 5 ==========
    heading(doc, "5. Development", 1)

    heading(doc, "5.1 Development Approach", 2)
    add_para(
        doc,
        "Development followed iterative vertical slices: foundation (project setup, auth, user model) → "
        "social feeds → shop/cart → Stripe orders → business subscription & dashboard → chat/notifications → "
        "AI assistant → admin console → i18n/polish. Frontend and backend advanced in parallel against agreed "
        "API contracts. Agile practices used: short cycles, demos, MoSCoW backlog refinement.",
    )
    caption(doc, "Table 6. Technology stack summary")
    add_table(
        doc,
        ["Layer", "Technologies"],
        [
            ["Client", "React 18, Vite, TypeScript, Tailwind CSS, Radix UI, Motion, Recharts, Axios, Socket.IO client, Stripe.js, Google OAuth, i18next"],
            ["Server", "Node.js, Express, Mongoose, JWT, bcrypt, Multer, Nodemailer, Morgan, CORS, Socket.IO, Stripe SDK, Google Auth Library"],
            ["Database", "MongoDB"],
            ["AI", "Python FastAPI disease classifier (Torch/Transformers); Hugging Face text assistance; local knowledge fallback; client/server moderation utilities"],
            ["Payments & other", "Stripe; Whish Money path; OpenStreetMap Nominatim / Lebanon locations; SMTP email"],
        ],
    )

    heading(doc, "5.2 Challenges and Problem-Solving", 2)
    bullet(doc, "Payment correctness (multi-leg seller/tax, webhooks vs client polling) — solved with PaymentIntent + webhook ledger writes.")
    bullet(doc, "Subscription gating of catalog visibility — solved with subscription status checks and shop filters.")
    bullet(doc, "Realtime support concurrency — solved with support conversation locks.")
    bullet(doc, "Bilingual RTL UX — solved with i18n resources and direction switching.")
    bullet(doc, "Practical AI without only cloud vision dependency — local Python disease service integrated into assistant flow.")
    bullet(doc, "Scope alignment of original “engineers” vision — realized via User community participation + AI advisory + business/admin roles.")
    add_para(
        doc,
        "Plan adjustments emphasized a robust web SPA for cross-device delivery within the course timeline, "
        "while preserving the original vision of social + commerce + intelligent support.",
    )

    heading(doc, "5.3 Testing and Debugging", 2)
    add_para(
        doc,
        "Testing combined manual/API scenario tests and integration checks for critical journeys: "
        "registration/login/Google, email verification, create post/thread, engage/save/follow, add-to-cart/"
        "checkout with Stripe test cards, subscription activate/renew (Stripe and Whish verify), chat send/"
        "receive, support lock behavior, admin report resolution, order mark-ready/cancel, AI text and photo "
        "paths, EN/AR switching. Defects were diagnosed via browser network/console, server logs, and Stripe "
        "test event traces. Performance attention included feed pagination, image handling, and dashboard refetch discipline.",
    )

    heading(doc, "5.4 Deployment and Version Control", 2)
    add_para(
        doc,
        "Local deployment runs client (Vite), API (Node), MongoDB, and optional Python AI service, with "
        "Stripe CLI listen for webhooks in development. Version control uses Git with a remote GitHub "
        "repository for collaboration and history. Environment files isolate secrets from source control. "
        "Production deployment targets static client build + Node host + managed MongoDB + configured "
        "webhook endpoints and environment secrets.",
    )

    heading(doc, "5.5 Teamwork and Collaboration", 2)
    add_para(
        doc,
        f"The project team consists of {TEAM[0]}, {TEAM[1]}, and {TEAM[2]}, under the supervision of "
        f"{SUPERVISOR}. Collaboration used GitHub for code and messaging channels for coordination. "
        "Cross-cutting concerns (auth, roles, i18n, payments) were reviewed jointly.",
    )
    add_para(doc, "Approximate contribution focus:")
    bullet(doc, f"{TEAM[0]} — frontend architecture, core UI pages, and client-side state/integration.")
    bullet(doc, f"{TEAM[1]} — backend APIs, data models, authentication/authorization, and server-side business logic.")
    bullet(doc, f"{TEAM[2]} — payments/subscriptions, AI assistant integration, testing support, and project documentation.")
    add_para(
        doc,
        "All three members participated in requirements workshops, UML design, demo rehearsals, and final "
        "end-to-end integration.",
    )

    heading(doc, "5.6 Reflection and Future Improvements", 2)
    add_para(
        doc,
        "The implemented platform meets the core expectations of the original vision: knowledge sharing, "
        "communication, online ordering, and intelligent support on one agricultural platform. Future "
        "improvements include richer recommendation ranking, expanded AI crop coverage, native mobile "
        "clients, deeper logistics integrations, formal expert consultation booking, and broader automated "
        "test coverage.",
    )

    page_break(doc)

    # ========== CHAPTER 6 ==========
    heading(doc, "6. Conclusion", 1)
    add_para(
        doc,
        "MASHTAL — Intelligent Agriculture Platform was specified, designed, planned, and implemented as a "
        "unified web system addressing real fragmentation in agricultural discovery, commerce, collaboration, "
        "and advisory support. Starting from a clear domain problem and product vision, the team produced "
        "structured requirements (including MoSCoW and use cases), a high-level architecture, UML use-case "
        "and class models, and a working full-stack implementation integrating social networking, e-commerce "
        "payments, business subscriptions, realtime messaging, administration, bilingual UX, and AI plant "
        "assistance. The project demonstrates applied software engineering practice across the full lifecycle "
        "expected by IN448.",
    )

    heading(doc, "6.1 Future Considerations", 2)
    bullet(doc, "Mobile applications and offline-friendly field workflows.")
    bullet(doc, "Expanded AI diagnostics and localized agronomy knowledge.")
    bullet(doc, "Advanced logistics and delivery partner integrations.")
    bullet(doc, "Stronger automated testing and observability in production.")
    bullet(doc, "Optional formal expert consultation marketplace features.")

    page_break(doc)

    # ========== BIBLIOGRAPHY ==========
    heading(doc, "Bibliography", 1)
    numbered(doc, "IEEE Computer Society. IEEE Recommended Practice for Software Requirements Specifications.")
    numbered(doc, "Object Management Group. OMG Unified Modeling Language (UML) specification.")
    numbered(doc, "React documentation. https://react.dev/")
    numbered(doc, "Express.js documentation. https://expressjs.com/")
    numbered(doc, "MongoDB Manual. https://www.mongodb.com/docs/")
    numbered(doc, "Stripe Documentation — Payments and Webhooks. https://stripe.com/docs")
    numbered(doc, "Socket.IO Documentation. https://socket.io/docs/")
    numbered(doc, "Vite Documentation. https://vitejs.dev/")
    numbered(doc, "i18next Documentation. https://www.i18next.com/")
    numbered(doc, "FastAPI Documentation. https://fastapi.tiangolo.com/")

    page_break(doc)

    # ========== APPENDICES ==========
    heading(doc, "Appendix A — Use Case Catalog Reference", 1)
    add_para(
        doc,
        "Figure 1 aggregates the platform use cases. For traceability, the implementation catalog groups "
        "capabilities as follows (non-exhaustive list of groups; each maps to UI pages and API routes):",
    )
    bullet(doc, "Guest: discover/home, posts/threads/shop browse, search/filters, product/business pages, AI, auth, legal pages.")
    bullet(doc, "User/Visitor: profile/media, follow/save, engage social, translate/share, cart/checkout, purchases, reviews/reports, chat/support, notifications.")
    bullet(doc, "Business inactive: renew (Stripe/Whish), profile, limited content; products hidden.")
    bullet(doc, "Business active: products CRUD, orders mark-ready, analytics KPIs/charts, posts/threads, customer chat.")
    bullet(doc, "Admin: overview, users, businesses, subscriptions, transactions, orders, reports, Whish verify, support locks, admin content.")
    bullet(doc, "External: Stripe, Google, SMTP, Whish, AI disease service, translate, locations, social share targets.")
    add_figure(doc, USECASE_IMG, 6.3, "Figure A1. Use-case diagram (repeated for appendix convenience)")

    heading(doc, "Appendix B — Class Diagram Reference", 1)
    add_para(
        doc,
        "Figure 3 / Figure B1 is the authoritative class diagram included in this report. It should be read "
        "together with Table 7 (model inventory) and Section 4.3 relationship highlights.",
    )
    add_figure(doc, CLASS_IMG, 6.3, "Figure B1. Class diagram (repeated for appendix convenience)")

    heading(doc, "Appendix C — Implementation Evidence Checklist", 1)
    add_para(doc, "Evaluators may verify the following implemented surfaces in a live or recorded demo:")
    bullet(doc, "Auth: email sign-up/in, Google OAuth, email verification, profile edit.")
    bullet(doc, "Social: posts & threads feeds; moderation blocks unsafe images/text; comments, likes, share, save, follow.")
    bullet(doc, "Search: unified search with filters.")
    bullet(doc, "Shop: catalog filters/sort, cart, Stripe checkout, purchase history, cancel/ready.")
    bullet(doc, "Business: subscription pay, dashboard analytics, product & order management.")
    bullet(doc, "Chat: DM + Mashtal Support; notifications center.")
    bullet(doc, "AI: text Q&A + plant photo analysis.")
    bullet(doc, "Admin: users/businesses/subscriptions/transactions/orders/reports/Whish verify.")
    bullet(doc, "i18n: EN/AR RTL switch.")

    heading(doc, "Appendix D — Stakeholder Value Overview", 1)
    caption(doc, "Figure 4. Conceptual stakeholder and value-flow overview")
    add_para(
        doc,
        "Users/Farmers gain knowledge, supplies, chat, and AI help. Businesses gain online storefront, "
        "orders, and analytics under subscription. Admins gain trust & finance control. External providers "
        "enable payments, identity, email, AI, and location. Value flows through content engagement, "
        "purchases (seller + platform fee legs), and subscription fees.",
        italic=True,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
